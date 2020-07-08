from docx import Document
from docx.shared import Cm
import os

tabela = (
    (1, 'Instalando Python', '3:05'),
    (2, 'Instalando VS Code', '5:50'),
    (3, 'Instalando Selenium', '3:50'),
)

word = Document()
word.add_heading('Olá, Mestres!', 0)
paragrafo = word.add_paragraph('Bem vindo a comunidade ')
paragrafo.add_run('Mestres da Automação ').bold = True
paragrafo.add_run('o lugar onde os mestres se encontram')
word.add_heading('Módulo 1 – Mestre do Software', level=1)
word.add_paragraph("Módulo 2 – Mestre do Bootcamp Python", style='Heading 2')
word.add_heading('Módulo 3 – Mestre da Web', level=3)
word.add_paragraph(
    '"Descubra como a programação pode mudar o futuro e a sua vida"', style="Quote")

word.add_paragraph('Python', style='List Bullet')
word.add_paragraph('Python3', style='List Bullet')
word.add_paragraph('pip', style='List Bullet')
word.add_paragraph('pip3', style='List Bullet')

word.add_paragraph('Windows', style='List Number')
word.add_paragraph('Linux', style='List Number')
word.add_paragraph('MacOs', style='List Number')


Table = word.add_table(rows=1, cols=3)

cabecalho = Table.rows[0].cells
cabecalho[0].text = 'Aula'
cabecalho[1].text = 'Nome'
cabecalho[2].text = 'Duração'

for aula, nome, duracao in tabela:
    dados_linha = Table.add_row().cells
    dados_linha[0].text = str(aula)
    dados_linha[1].text = nome
    dados_linha[2].text = duracao

word.add_picture('Módulo 4'+os.sep+'word'+os.sep+'note.jpg', width=Cm(5.25))
word.save('Módulo 4'+os.sep+'word'+os.sep+'demo.docx')
